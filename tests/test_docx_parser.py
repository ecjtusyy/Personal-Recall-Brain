from docx import Document
from docx.shared import Inches
from PIL import Image

from second_brain.ingestion.docx_parser import parse_docx


def test_docx_preserves_paragraph_and_table_order(tmp_path):
    path = tmp_path / "测试.docx"
    doc = Document()
    doc.add_heading("申论复习", level=1)
    doc.add_paragraph("第一段")
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "公文写作"
    table.cell(0, 1).text = "主体分析"
    doc.add_paragraph("")
    doc.add_paragraph("最后一段")
    doc.save(path)
    parsed = parse_docx(path)
    texts = [block.text for block in parsed.text_blocks]
    assert texts == ["申论复习", "第一段", "公文写作 | 主体分析", "最后一段"]


def test_docx_records_embedded_image(tmp_path):
    image_path = tmp_path / "note.png"
    Image.new("RGB", (32, 32), "white").save(image_path)
    path = tmp_path / "含图片.docx"
    doc = Document()
    doc.add_paragraph("图片前")
    doc.add_picture(str(image_path), width=Inches(1))
    doc.save(path)
    parsed = parse_docx(path)
    assert len(parsed.assets) == 1
    assert parsed.assets[0].original_name.endswith(".png")
    assert parsed.assets[0].data

